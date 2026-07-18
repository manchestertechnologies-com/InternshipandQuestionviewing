import docx
import os
import json

def inspect_file(filename):
    print(f"=== Inspecting {filename} ===")
    if not os.path.exists(filename):
        print(f"File {filename} not found!")
        return
    doc = docx.Document(filename)
    print(f"Number of tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"Table {i} has {len(table.rows)} rows and {len(table.columns)} columns.")
        # Print headers (row 0)
        if len(table.rows) > 0:
            headers = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
            print(f"Headers: {headers}")
            # Print first 2 data rows
            for r_idx in range(1, min(3, len(table.rows))):
                row_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[r_idx].cells]
                print(f"Row {r_idx}: {row_cells}")

inspect_file("PALLAVI_DS_updated.docx")
inspect_file("PRATEEKSHA_updated.docx")
inspect_file("RAMYA_S_updated.docx")
